from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT_FILE = DOCS_DIR / "informe_tecnico.docx"


def add_heading_center(document: Document, text: str, level: int = 0) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    add_heading_center(doc, "Informe técnico del caso práctico N.º 01", 0)
    add_heading_center(doc, "Business Corporation - Sistema de Recursos Humanos", 1)

    p = doc.add_paragraph()
    p.add_run("Fecha de generación: ").bold = True
    p.add_run(str(date.today()))

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Se desarrolló una solución en Python con interfaz Streamlit para gestionar la estructura del área de Recursos Humanos de Business Corporation. "
        "La solución aplica programación orientada a objetos, lista de objetos, herencia, encapsulamiento, polimorfismo, condicionales y estructuras repetitivas."
    )

    doc.add_heading("2. Requisitos interpretados del caso", level=1)
    add_bullet_list(
        doc,
        [
            "Implementar un array de objetos con trabajadores.",
            "Registrar como mínimo un gerente y cinco jefes de área.",
            "Registrar asistentes y personal técnico subordinados a cada jefe de área.",
            "Controlar que un área no supere dos asistentes ni cinco técnicos.",
            "Mostrar nombre completo, resumen del trabajador, jefe inmediato y estado.",
            "Implementar getters y setters para mantenimiento del sistema.",
        ],
    )

    doc.add_heading("3. Criterio de reconstrucción", level=1)
    doc.add_paragraph(
        "Esta versión fue rehecha para cumplir la indicación de trabajar todo en Python. "
        "Por ello, se eliminaron las tablas HTML manuales, el flujograma en formatos externos y el script de entorno en shell."
    )

    doc.add_heading("4. Arquitectura propuesta", level=1)
    doc.add_paragraph(
        "Se diseñó una arquitectura modular con separación entre modelos, servicios, utilidades, pruebas y documentación. "
        "Esto reduce acoplamiento, facilita mantenimiento y vuelve más sólida la defensa académica del proyecto."
    )

    doc.add_heading("5. Explicación archivo por archivo", level=1)
    explicaciones = {
        "app.py": "Contiene la interfaz Streamlit y usa componentes nativos para mostrar la información, sin HTML manual.",
        "models/trabajador.py": "Define la clase abstracta base con atributos encapsulados, getters, setters y métodos comunes.",
        "models/roles.py": "Declara las subclases Gerente, JefeArea, Asistente y PersonalTecnico, aplicando polimorfismo en get_resumen().",
        "services/repository.py": "Gestiona la colección de trabajadores y valida reglas del negocio.",
        "services/data_factory.py": "Carga los datos demo exigidos por la consigna, respetando jerarquías y límites por área.",
        "services/reporting.py": "Convierte los objetos en DataFrames de Pandas para visualización y análisis dentro de Streamlit.",
        "utils/constants.py": "Centraliza constantes de estados y áreas.",
        "utils/helpers.py": "Incluye funciones auxiliares y decisiones documentadas del diseño.",
        "tests/test_logic.py": "Pruebas automáticas para verificar el cumplimiento de reglas clave.",
        "scripts/setup_env.py": "Automatiza en Python la creación del entorno virtual e instalación de dependencias.",
        "docs/flujograma_pseudocodigo.txt": "Describe el flujo del sistema en texto estructurado y pseudocódigo.",
        "README.md": "Documento guía del repositorio con instalación, ejecución, estructura y enlace futuro del repositorio.",
    }
    for archivo, descripcion in explicaciones.items():
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{archivo}: ").bold = True
        paragraph.add_run(descripcion)

    doc.add_heading("6. Evidencia de cumplimiento técnico", level=1)
    add_bullet_list(
        doc,
        [
            "POO: uso de clase abstracta y clases derivadas.",
            "Herencia: especialización de trabajador por rol.",
            "Encapsulamiento: atributos protegidos mediante getters y setters.",
            "Polimorfismo: redefinición de get_resumen() por tipo de trabajador.",
            "If y match-case: control de estados y decisiones lógicas.",
            "Interfaz resuelta con Streamlit usando tablas nativas basadas en DataFrames.",
        ],
    )

    doc.add_heading("7. Entorno virtual y ejecución", level=1)
    doc.add_paragraph("Comandos principales:")
    for command in [
        "python scripts/setup_env.py",
        ".venv/bin/streamlit run app.py",
        ".venv/bin/pytest -q",
        ".venv/bin/python generate_report.py",
    ]:
        doc.add_paragraph(command, style="List Bullet")

    doc.add_heading("8. Repositorio", level=1)
    doc.add_paragraph(
        "La carpeta queda preparada para control de versiones con Git. Si se publica en GitHub o GitLab, el enlace debe colocarse en el README y en esta sección."
    )
    doc.add_paragraph("Enlace actual: PENDIENTE DE PUBLICACIÓN REMOTA")

    doc.add_heading("9. Observación crítica", level=1)
    doc.add_paragraph(
        "El enunciado presenta una inconsistencia: enumera cuatro áreas específicas bajo gerencia, pero exige cinco jefes de área. "
        "Se resolvió añadiendo el área de Finanzas para cumplir el mínimo cuantitativo solicitado."
    )

    doc.add_heading("10. Conclusión", level=1)
    doc.add_paragraph(
        "La solución entregada es funcional, modular, ampliable y más coherente con la instrucción de trabajar todo en Python. "
        "No se apoya en HTML manual ni en artificios de presentación externos para resolver el caso."
    )

    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    main()
